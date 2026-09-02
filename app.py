import io
import json
import os
import re
import time
from typing import Any

import streamlit as st
from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document

# OCR dependencies for scanned/image-only PDFs
import fitz  # PyMuPDF
import pytesseract
from PIL import Image, ImageOps

APP_TITLE = "ResumeForge AI — ATS Resume Analyzer"
MODEL_NAME = "gemini-3.6-flash"
MAX_TEXT_CHARS = 50000

st.set_page_config(page_title=APP_TITLE, page_icon="📄", layout="wide")

st.markdown("""
<style>
.block-container {max-width: 1180px; padding-top: 2rem; padding-bottom: 3rem;}
.hero {padding: 1.5rem 1.6rem; border-radius: 18px; border: 1px solid rgba(128,128,128,.25); margin-bottom: 1rem;}
.small {color: #777; font-size: .92rem;}
.score {font-size: 3.4rem; font-weight: 800; line-height: 1;}
.card {padding: 1rem; border-radius: 14px; border: 1px solid rgba(128,128,128,.22); margin-bottom: .8rem;}
</style>
""", unsafe_allow_html=True)


def get_api_key() -> str | None:
    key = None
    try:
        key = st.secrets.get("GEMINI_API_KEY")
    except Exception:
        pass
    return key or os.getenv("GEMINI_API_KEY")


def ocr_pdf_to_text(data: bytes) -> str:
    """Extract text from an image-only/scanned PDF using Tesseract OCR."""
    pages_text = []

    try:
        pdf = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise ValueError(f"Could not open the PDF: {exc}") from exc

    try:
        # Resumes are normally short, so limit OCR work to 10 pages.
        for page_number in range(min(len(pdf), 10)):
            page = pdf.load_page(page_number)
            pix = page.get_pixmap(dpi=200, alpha=False)
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Basic preprocessing improves OCR on many scanned resumes.
            gray = ImageOps.grayscale(image)
            gray = ImageOps.autocontrast(gray)

            page_text = pytesseract.image_to_string(
                gray,
                lang="eng",
                config="--oem 3 --psm 6",
            )

            if page_text.strip():
                pages_text.append(page_text.strip())
    finally:
        pdf.close()

    return "\n\n".join(pages_text).strip()


def extract_resume_text(uploaded_file) -> tuple[str, bool]:
    """
    Extract resume text.
    Returns: (text, used_ocr)
    """
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()
    used_ocr = False

    if name.endswith(".pdf"):
        # First try normal PDF text extraction.
        reader = PdfReader(io.BytesIO(data))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)

        # Automatically fall back to OCR for scanned/image-only PDFs.
        if len(re.sub(r"\s+", "", text)) < 80:
            used_ocr = True
            text = ocr_pdf_to_text(data)

    elif name.endswith(".docx"):
        doc = Document(io.BytesIO(data))
        chunks = [p.text for p in doc.paragraphs]

        for table in doc.tables:
            for row in table.rows:
                chunks.append(" | ".join(cell.text for cell in row.cells))

        text = "\n".join(chunks)

    elif name.endswith(".txt"):
        text = data.decode("utf-8", errors="ignore")

    else:
        raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")

    cleaned = re.sub(r"\n{3,}", "\n\n", text).strip()[:MAX_TEXT_CHARS]

    if not cleaned:
        if used_ocr:
            raise ValueError(
                "OCR could not read meaningful text from this PDF. "
                "Please upload a clearer PDF or a DOCX/text-based resume."
            )
        raise ValueError(
            "No readable text was found. Please upload a text-based PDF, DOCX, or TXT resume."
        )

    return cleaned, used_ocr


def deterministic_ats_checks(text: str) -> dict[str, Any]:
    lower = text.lower()
    checks = {
        "contact": bool(re.search(r"(?:email|@)|(?:\+?\d[\d\s().-]{7,})", lower)),
        "summary": bool(re.search(r"\b(summary|profile|objective|professional summary)\b", lower)),
        "experience": bool(re.search(r"\b(experience|employment|work history|professional experience)\b", lower)),
        "education": bool(re.search(r"\b(education|academic background|qualifications)\b", lower)),
        "skills": bool(re.search(r"\b(skills|technical skills|core competencies|competencies)\b", lower)),
        "achievements": bool(re.search(r"\b(achievements|accomplishments|awards|projects)\b", lower)),
        "action_verbs": len(re.findall(r"\b(achieved|built|created|developed|designed|improved|increased|reduced|managed|led|analyzed|implemented|delivered|optimized|automated)\b", lower)) >= 3,
        "quantification": bool(re.search(r"\b\d+(?:\.\d+)?\s*(?:%|percent|k|m|million|thousand|years?|months?)\b|\$\s*\d+", lower)),
    }
    score = round(sum(checks.values()) / len(checks) * 30)
    return {"checks": checks, "baseline_score": score}


def extract_json(text: str) -> dict[str, Any]:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", cleaned, re.S)
        if not match:
            raise ValueError("Gemini returned an unreadable analysis. Please try again.")
        return json.loads(match.group(0))


def analyze_with_gemini(resume_text: str, job_description: str, baseline: dict[str, Any]) -> dict[str, Any]:
    api_key = get_api_key()
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY is missing. Add it to Streamlit Secrets or your environment.")

    client = genai.Client(api_key=api_key)
    jd = job_description.strip() or "No job description supplied. Evaluate general ATS readiness and employer appeal."
    prompt = f"""
You are a senior ATS optimization specialist, recruiter, and executive resume writer.
Analyze the candidate resume below. Be rigorous, evidence-based, and never invent experience,
metrics, skills, employers, degrees, certifications, or achievements that are not supported by the resume.
The goal is to maximize ATS compatibility AND make the candidate more persuasive to a real employer.

IMPORTANT ATS RULE:
An ATS score is an estimate, not a universal official score. Explain that it is a practical readiness
score based on structure, parsing, relevance, keywords, and evidence.

Return ONLY valid JSON matching the requested schema.

JOB DESCRIPTION:
{jd[:25000]}

DETERMINISTIC CHECKS:
{json.dumps(baseline, ensure_ascii=False)}

RESUME:
{resume_text}

JSON schema:
{{
  "ats_score": 0,
  "ats_band": "Needs work | Fair | Strong | Excellent",
  "executive_verdict": "short recruiter-style verdict",
  "summary": "2-4 sentence assessment",
  "category_scores": {{
    "formatting_and_parsing": 0,
    "section_structure": 0,
    "keyword_alignment": 0,
    "experience_impact": 0,
    "skills_and_relevance": 0,
    "employer_persuasion": 0
  }},
  "critical_issues": ["..."],
  "quick_wins": ["..."],
  "missing_keywords": ["keyword or phrase actually present in JD and relevant to resume"],
  "keyword_matches": ["relevant matched keyword/phrase"],
  "bullet_rewrites": [
    {{"original": "existing resume bullet", "improved": "stronger version without inventing facts", "why": "reason"}}
  ],
  "professional_summary": "rewritten summary using only supported facts",
  "ats_format_recommendations": ["..."],
  "employer_persuasion_recommendations": ["..."],
  "final_action_plan": ["priority 1", "priority 2", "priority 3", "priority 4", "priority 5"]
}}

Scoring guidance: 0-100. Consider the deterministic checks but independently judge the complete resume.
If no job description is provided, score general ATS readiness and do not claim keyword alignment with a specific role.
"""
        response = client.models.generate_content(
        model=MODEL_NAME,
        contents=prompt,
        config=types.GenerateContentConfig(
            temperature=0.2,
            response_mime_type="application/json",
        ),
    ) 

    result = extract_json(response.text)
    result["ats_score"] = max(0, min(100, int(result.get("ats_score", 0))))
    return result


def show_list(title: str, items: list[str], empty: str = "Nothing major detected."):
    st.subheader(title)
    if not items:
        st.info(empty)
    for item in items:
        st.markdown(f"- {item}")


st.markdown(f"<div class='hero'><h1>📄 {APP_TITLE}</h1><p>Upload a resume → get an estimated ATS score → fix weaknesses → make the resume stronger and more persuasive.</p><p class='small'>Powered by Gemini 2.5 Flash. No fabricated claims: recommendations are grounded in the uploaded resume.</p></div>", unsafe_allow_html=True)

with st.sidebar:
    st.header("⚙️ Analysis settings")
    st.caption("Gemini model")
    st.code(MODEL_NAME)
    st.caption("Tip: add a job description for targeted keyword matching.")
    st.divider()
    st.markdown("**What you get**")
    st.markdown("✓ ATS readiness score\n\n✓ Category breakdown\n\n✓ Missing/relevant keywords\n\n✓ Critical issues & quick wins\n\n✓ Bullet rewrites\n\n✓ Recruiter-style summary\n\n✓ Employer persuasion advice")

uploaded = st.file_uploader("Upload your resume", type=["pdf", "docx", "txt"], help="PDF, DOCX, or TXT. Avoid uploading confidential information you do not want processed by a third-party AI service.")
job_description = st.text_area("Optional: paste the job description", height=220, placeholder="Paste the target job description here for role-specific ATS keyword matching...")

if uploaded:
    try:
        resume_text, used_ocr = extract_resume_text(uploaded)
    except Exception as exc:
        st.error(f"Could not read the resume: {exc}")
        st.stop()

    if used_ocr:
        st.success("Scanned/image-based PDF detected — OCR was automatically used to read the resume.")

    if len(resume_text.strip()) < 80:
        st.warning("Only a small amount of text could be extracted. For best results, upload a clearer resume.")

    with st.expander("Preview extracted resume text"):
        st.text(resume_text[:10000])

    if st.button("🚀 Analyze & Improve My Resume", type="primary", use_container_width=True):
        with st.spinner("Analyzing structure, ATS readiness, keywords, impact, and recruiter appeal..."):
            try:
                baseline = deterministic_ats_checks(resume_text)
                result = analyze_with_gemini(resume_text, job_description, baseline)
                st.session_state["analysis"] = result
            except Exception as exc:
                st.error(f"Analysis failed: {exc}")
                st.info("Check that your Gemini API key is valid and that the selected Gemini model is available to your API project.")

analysis = st.session_state.get("analysis")
if analysis:
    st.divider()
    score = analysis.get("ats_score", 0)
    band = analysis.get("ats_band", "")
    c1, c2, c3 = st.columns([1, 1.3, 2])
    with c1:
        st.markdown(f"<div class='score'>{score}/100</div><div>Estimated ATS readiness</div>", unsafe_allow_html=True)
    with c2:
        st.metric("ATS band", band)
    with c3:
        st.info(analysis.get("executive_verdict", ""))

    st.caption("This is an AI-assisted estimate, not a score produced by every ATS in the market. Different employers use different systems and configurations.")
    st.write(analysis.get("summary", ""))

    st.subheader("📊 Score breakdown")
    category_scores = analysis.get("category_scores", {})
    cols = st.columns(3)
    for i, (name, value) in enumerate(category_scores.items()):
        cols[i % 3].metric(name.replace("_", " ").title(), f"{int(value)}/100")

    left, right = st.columns(2)
    with left:
        show_list("🚨 Critical issues", analysis.get("critical_issues", []))
        show_list("⚡ Quick wins", analysis.get("quick_wins", []))
    with right:
        show_list("🔎 Missing keywords", analysis.get("missing_keywords", []), "No major missing keywords identified.")
        show_list("✅ Matched keywords", analysis.get("keyword_matches", []), "No targeted keyword matching was available.")

    st.subheader("✍️ Stronger professional summary")
    st.success(analysis.get("professional_summary", ""))

    st.subheader("🧲 Make your experience more persuasive")
    rewrites = analysis.get("bullet_rewrites", [])
    if rewrites:
        for item in rewrites:
            with st.expander(item.get("original", "Resume bullet")):
                st.markdown("**Improved version**")
                st.write(item.get("improved", ""))
                st.caption(item.get("why", ""))
    else:
        st.info("No bullet rewrites were returned.")

    show_list("🧾 ATS formatting recommendations", analysis.get("ats_format_recommendations", []))
    show_list("💼 Employer-persuasion recommendations", analysis.get("employer_persuasion_recommendations", []))
    show_list("🎯 Final action plan", analysis.get("final_action_plan", []))

    st.download_button(
        "⬇️ Download analysis as JSON",
        data=json.dumps(analysis, indent=2, ensure_ascii=False),
        file_name="resume_ats_analysis.json",
        mime="application/json",
    )
else:
    st.info("Upload a resume and click **Analyze & Improve My Resume** to begin.")
